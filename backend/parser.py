import io
import re
import sys
from typing import Optional

import docx        # python-docx  — Word (.docx)
import fitz        # PyMuPDF      — PDF

_GITHUB_RE = re.compile(
    r"https?://(?:www\.)?github\.com/[A-Za-z0-9_.-]+(?:/[A-Za-z0-9_.-]+)?",
    re.IGNORECASE,
)

_EMAIL_RE = re.compile(
    r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}",
    re.IGNORECASE,
)

_PDF_MAGIC  = b"%PDF"
_DOCX_MAGIC = b"PK\x03\x04"   # DOCX is a ZIP archive


def _normalize(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = "\n".join(re.sub(r"[ \t]+", " ", line).strip() for line in text.split("\n"))
    return text.strip()


def _unique_ordered(items: list[str]) -> list[str]:
    seen: set[str] = set()
    result: list[str] = []
    for item in items:
        key = item.lower()
        if key not in seen:
            seen.add(key)
            result.append(item)
    return result


# ---------------------------------------------------------------------------
# PDF extraction
# ---------------------------------------------------------------------------

def _extract_pdf(file_bytes: bytes) -> tuple[str, int]:
    try:
        doc = fitz.open(stream=io.BytesIO(file_bytes), filetype="pdf")
    except fitz.FileDataError as exc:
        raise ValueError(f"Not a valid PDF file: {exc}") from exc
    except Exception as exc:
        raise ValueError(f"Failed to open PDF: {exc}") from exc

    try:
        if doc.is_encrypted:
            if doc.authenticate("") == 0:
                raise ValueError(
                    "PDF is password-protected and could not be decrypted. "
                    "Please provide an unlocked copy."
                )
        page_count = doc.page_count
        if page_count == 0:
            raise ValueError("PDF has no pages.")

        page_texts: list[str] = []
        for page in doc:
            try:
                page_texts.append(page.get_text())
            except Exception as exc:
                page_texts.append("")
                print(f"[parser] Warning: could not read page {page.number}: {exc}", file=sys.stderr)
    finally:
        doc.close()

    return "\n\n".join(page_texts), page_count


# ---------------------------------------------------------------------------
# DOCX extraction
# ---------------------------------------------------------------------------

def _extract_docx(file_bytes: bytes) -> tuple[str, int]:
    try:
        document = docx.Document(io.BytesIO(file_bytes))
    except Exception as exc:
        raise ValueError(f"Not a valid Word (.docx) file: {exc}") from exc

    parts: list[str] = []

    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)

    # Pull hyperlinks from document relationships — catches GitHub links
    # that are anchor text rather than raw URLs in the body.
    try:
        for rel in document.part.rels.values():
            if "hyperlink" in rel.reltype and rel.is_external:
                target = rel.target_ref
                if target and target.startswith("http"):
                    parts.append(target)
    except Exception:
        pass

    raw_text = "\n".join(parts)
    estimated_pages = max(1, round(len(raw_text.split()) / 300))
    return raw_text, estimated_pages


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def extract_resume_text(file_bytes: bytes, filename: str = "") -> dict:
    """
    Extract structured text from a PDF or Word (.docx) résumé supplied as bytes.

    File type is determined from magic bytes — not the filename — so renamed
    files still parse correctly.

    Returns:
        full_text, page_count, github_urls, email, word_count, file_type
    Raises:
        ValueError on empty input, unsupported format, or unreadable content.
    """
    if not file_bytes:
        raise ValueError("No file bytes provided — cannot parse an empty payload.")

    if file_bytes[:4] == _PDF_MAGIC or file_bytes[:5] == b"%PDF-":
        raw_text, page_count = _extract_pdf(file_bytes)
        file_type = "pdf"
    elif file_bytes[:4] == _DOCX_MAGIC:
        raw_text, page_count = _extract_docx(file_bytes)
        file_type = "docx"
    else:
        label = f" '{filename}'" if filename else ""
        raise ValueError(
            f"Unsupported file type{label}. Only PDF (.pdf) and Word (.docx) files are accepted."
        )

    full_text = _normalize(raw_text)

    if not full_text:
        raise ValueError(
            "The file contains no extractable text. "
            "If it is a scanned image, OCR is required."
        )

    github_urls = _unique_ordered(_GITHUB_RE.findall(full_text))
    email_match = _EMAIL_RE.search(full_text)
    email: Optional[str] = email_match.group(0).lower() if email_match else None

    return {
        "full_text":   full_text,
        "page_count":  page_count,
        "github_urls": github_urls,
        "email":       email,
        "word_count":  len(full_text.split()),
        "file_type":   file_type,
    }


# ---------------------------------------------------------------------------
# Quick test:  python parser.py resume.pdf   or   python parser.py resume.docx
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python parser.py <resume.pdf|resume.docx>")
        sys.exit(1)

    try:
        with open(sys.argv[1], "rb") as f:
            raw = f.read()
        result = extract_resume_text(raw, filename=sys.argv[1])
    except (FileNotFoundError, ValueError) as e:
        print(f"Error: {e}")
        sys.exit(1)

    print(f"Type       : {result['file_type'].upper()}")
    print(f"Pages      : {result['page_count']}")
    print(f"Words      : {result['word_count']}")
    print(f"Email      : {result['email'] or '(none found)'}")
    print(f"GitHub URLs: {result['github_urls'] or '(none found)'}")
    print("\n--- Full Text (first 500 chars) ---")
    print(result["full_text"][:500])
